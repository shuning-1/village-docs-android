#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""智慧文档管理系统 - Flask后端服务"""

import os
import sys
import json
import uuid
import glob
import re
import tempfile
import hmac
import secrets
from datetime import datetime
from flask import Flask, request, jsonify, send_from_directory
import base64
import hashlib
import openpyxl
import xlrd
import csv
import calendar

app = Flask(__name__, static_folder=None)
app.config['MAX_CONTENT_LENGTH'] = 12 * 1024 * 1024

# ===== 路径配置 =====
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
# PyInstaller frozen 模式：使用 exe 所在目录（即安装目录）作为数据目录
# 所有数据（配置、上传文件等）都存放在用户安装的文件夹里
if getattr(sys, 'frozen', False):
    _data_dir = os.environ.get('_data_dir', '')
    if not _data_dir:
        _data_dir = os.path.dirname(sys.executable)
    BASE_DIR = _data_dir
CONFIG_PATH = os.path.join(BASE_DIR, 'config.json')
LOGIN_CONFIG_PATH = os.path.join(BASE_DIR, 'login_config.json')
DASHBOARD_CONFIG_PATH = os.path.join(BASE_DIR, 'dashboard_config.json')
CERT_DATA_PATH = os.path.join(BASE_DIR, 'cert_data.json')
DISMISSED_REMINDERS_PATH = os.path.join(BASE_DIR, 'dismissed_reminders.json')
ACCOUNTS_PATH = os.path.join(BASE_DIR, 'accounts.json')
UPLOADS_DIR = os.path.join(BASE_DIR, 'uploads')

DOCUMENT_EXTENSIONS = ('.xls', '.xlsx', '.csv', '.doc', '.docx', '.pdf')
EXCEL_EXTENSIONS = ('.xls', '.xlsx', '.csv')
ALLOWED_IMAGE_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif', 'webp'}
MAX_IMAGE_BYTES = 5 * 1024 * 1024
PASSWORD_HASH_PREFIX = 'pbkdf2_sha256$'

DEFAULT_CONFIG = {
    "village_name": "",
    "town_name": "",
    "county_name": "",
    "base_path": BASE_DIR,
    "categories": []
}

DEFAULT_LOGIN_APPEARANCE = {
    "login_title": "智慧文档管理系统",
    "login_subtitle": "党建引领智慧村务平台",
    "login_icon_type": "emoji",
    "login_icon_emoji": "⭐",
    "login_icon_image": "",
    "login_icon_crop": {"x": 0, "y": 0, "w": 100, "h": 100, "size": 80},
    "login_bg_type": "color",
    "login_bg_color": "linear-gradient(135deg, #8B0000 0%, #C41E3A 40%, #A01030 100%)",
    "login_bg_image": "",
    "login_bg_fit": "cover",
    "header_title": "",
    "header_subtitle": "",
    "login_footer_text": "党建引领 · 智慧村务",
}


# ===== 配置读写 =====
def load_json_file(path, default):
    """安全加载 JSON，文件不存在或损坏时返回默认值副本"""
    if not os.path.exists(path):
        return json.loads(json.dumps(default, ensure_ascii=False))
    try:
        with open(path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        return data if data is not None else json.loads(json.dumps(default, ensure_ascii=False))
    except (json.JSONDecodeError, OSError):
        return json.loads(json.dumps(default, ensure_ascii=False))


def save_json_file(path, data):
    """原子化写入 JSON，避免程序异常退出导致配置文件写坏"""
    os.makedirs(os.path.dirname(path) or '.', exist_ok=True)
    fd, tmp_path = tempfile.mkstemp(prefix='.tmp_', suffix='.json', dir=os.path.dirname(path) or '.')
    try:
        with os.fdopen(fd, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
            f.write('\n')
        os.replace(tmp_path, path)
    finally:
        if os.path.exists(tmp_path):
            try:
                os.remove(tmp_path)
            except OSError:
                pass


def sanitize_filename(filename):
    """只允许使用文件名部分，阻止 ../ 形式的路径穿越"""
    filename = os.path.basename(str(filename or '').replace('\\', '/')).strip()
    if not filename or filename in {'.', '..'} or '\x00' in filename:
        raise ValueError('文件名无效')
    return filename


def resolve_base_file(base_path, filename, must_exist=False):
    """解析资料目录内的文件路径，确保最终路径仍在资料目录中"""
    filename = sanitize_filename(filename)
    base_abs = os.path.abspath(base_path or BASE_DIR)
    file_abs = os.path.abspath(os.path.join(base_abs, filename))
    if os.path.commonpath([base_abs, file_abs]) != base_abs:
        raise ValueError('非法文件路径')
    if must_exist and not os.path.exists(file_abs):
        raise FileNotFoundError('文件不存在')
    return file_abs


def hash_password(password):
    salt = secrets.token_hex(16)
    digest = hashlib.pbkdf2_hmac('sha256', str(password).encode('utf-8'), salt.encode('utf-8'), 120000)
    return f'{PASSWORD_HASH_PREFIX}{salt}${digest.hex()}'


def verify_password(account, password):
    """兼容旧版明文密码，优先验证哈希"""
    password = str(password or '')
    stored_hash = account.get('password_hash', '')
    if stored_hash.startswith(PASSWORD_HASH_PREFIX):
        try:
            _, salt, digest = stored_hash.split('$', 2)
            calc = hashlib.pbkdf2_hmac('sha256', password.encode('utf-8'), salt.encode('utf-8'), 120000).hex()
            return hmac.compare_digest(calc, digest)
        except ValueError:
            return False
    return hmac.compare_digest(str(account.get('password', '')), password)


def migrate_account_password(account):
    """登录或保存账号时把旧版明文密码迁移为哈希"""
    if account.get('password_hash', '').startswith(PASSWORD_HASH_PREFIX):
        account.pop('password', None)
        return False
    plain = account.get('password')
    if plain is None:
        return False
    account['password_hash'] = hash_password(plain)
    account.pop('password', None)
    return True


def load_config():
    default = dict(DEFAULT_CONFIG)
    return load_json_file(CONFIG_PATH, default)


def save_config(cfg):
    save_json_file(CONFIG_PATH, cfg)


def load_login_config():
    default = {
        "admin_user": "admin", "admin_pass": "123456", "secret_key": "admin",
        "login_village": "", "remembered_user": "", "remembered_pass": "",
        "login_title": "智慧文档管理系统", "login_subtitle": "党建引领智慧村务平台",
        "login_icon_type": "emoji",   # "emoji" | "image"
        "login_icon_emoji": "⭐",
        "login_icon_image": "",        # base64 data URL 或 /uploads/xxx 路径
        "login_icon_crop": {"x": 0, "y": 0, "w": 100, "h": 100, "size": 80},
        "login_bg_type": "color",      # "color" | "image"
        "login_bg_color": "linear-gradient(135deg, #8B0000 0%, #C41E3A 40%, #A01030 100%)",
        "login_bg_image": "",          # base64 data URL 或 /uploads/xxx 路径
        "login_bg_fit": "cover",       # "cover" | "contain" | "fill"
        "header_title": "",            # 顶部标题，空则使用 config 里的村名
        "header_subtitle": "",         # 顶部副标题
        "login_footer_text": "党建引领 · 智慧村务",   # 登录页底部文字
        "window_title": "智慧文档管理系统",            # EXE窗口标题
    }
    data = load_json_file(LOGIN_CONFIG_PATH, default)
    for k, v in default.items():
        if k not in data:
            data[k] = v
    return data


def load_accounts():
    """加载账户列表（多用户体系）"""
    default_password = hash_password("123456")
    default = {
        "accounts": [
            {
                "id": "super_1",
                "username": "admin",
                "password_hash": default_password,
                "role": "super",          # super / admin / staff
                "display_name": "超级管理员",
                "created_at": "",
                "enabled": True
            }
        ]
    }
    data = load_json_file(ACCOUNTS_PATH, default)
    if isinstance(data, list):
        data = {"accounts": data}
    return data


def save_accounts(data):
    if isinstance(data, list):
        data = {"accounts": data}
    for account in data.get('accounts', []):
        migrate_account_password(account)
    save_json_file(ACCOUNTS_PATH, data)




def save_login_config(cfg):
    save_json_file(LOGIN_CONFIG_PATH, cfg)


def load_dashboard_config():
    """加载看板自定义配置"""
    # 默认看板卡片
    default = {
        "cards": [
            {"id": "dc_1", "name": "在册村人数", "icon": "👥", "color": "#e74c3c", "file": "", "calc_type": "auto", "manual_value": 0},
            {"id": "dc_2", "name": "党员人数", "icon": "🚩", "color": "#c0392b", "file": "", "calc_type": "auto", "manual_value": 0},
            {"id": "dc_3", "name": "18岁以上人数", "icon": "🧑", "color": "#3498db", "file": "", "calc_type": "auto", "manual_value": 0},
            {"id": "dc_4", "name": "18岁以下人数", "icon": "👦", "color": "#f39c12", "file": "", "calc_type": "auto", "manual_value": 0},
            {"id": "dc_5", "name": "60岁以上人数", "icon": "🧓", "color": "#27ae60", "file": "", "calc_type": "auto", "manual_value": 0},
            {"id": "dc_6", "name": "80岁以上人数", "icon": "👴", "color": "#9b59b6", "file": "", "calc_type": "auto", "manual_value": 0},
        ],
        "analysis_file": "",
        "tab_order": ["files", "dashboard", "search", "reminders", "templates"]
    }
    return load_json_file(DASHBOARD_CONFIG_PATH, default)


def save_dashboard_config(cfg):
    save_json_file(DASHBOARD_CONFIG_PATH, cfg)


def load_cert_data():
    """加载证明数据"""
    return load_json_file(CERT_DATA_PATH, {"certificates": [], "next_id": 1})


def save_cert_data(data):
    save_json_file(CERT_DATA_PATH, data)


def calc_age_from_idcard(id_str):
    """从身份证号码计算年龄"""
    id_str = str(id_str).strip()
    if len(id_str) == 18:
        try:
            birth_year = int(id_str[6:10])
            birth_month = int(id_str[10:12])
            birth_day = int(id_str[12:14])
            today = datetime.now()
            age = today.year - birth_year
            if (today.month, today.day) < (birth_month, birth_day):
                age -= 1
            return age
        except ValueError:
            return None
    elif len(id_str) == 15:
        try:
            birth_year = int('19' + id_str[6:8])
            birth_month = int(id_str[8:10])
            birth_day = int(id_str[10:12])
            today = datetime.now()
            age = today.year - birth_year
            if (today.month, today.day) < (birth_month, birth_day):
                age -= 1
            return age
        except ValueError:
            return None
    return None


def auto_detect_calc_type(name):
    """根据卡片名称自动判断计算方式"""
    name = name.strip()
    if '80岁以上' in name or '80周岁以上' in name:
        return 'age_ge_80'
    if '60岁以上' in name or '60周岁以上' in name:
        return 'age_ge_60'
    if '18岁以上' in name or '18周岁以上' in name:
        return 'age_ge_18'
    if '18岁以下' in name or '18周岁以下' in name:
        return 'age_lt_18'
    if '人数' in name or '人口' in name:
        return 'count_rows'
    if '党员' in name:
        return 'count_rows'
    return 'count_rows'


def calculate_card_value(card, base_path):
    """计算看板卡片的数值"""
    calc_type = card.get('calc_type', 'auto')
    # 手动输入直接返回设定值
    if calc_type == 'manual':
        return card.get('manual_value', 0)

    filename = card.get('file', '')
    if not filename:
        return card.get('manual_value', 0) or 0

    try:
        filepath = resolve_base_file(base_path, filename, must_exist=True)
    except (ValueError, FileNotFoundError):
        return 0

    sheets, _, err = read_excel_file(filepath)
    if err or not sheets:
        return 0

    sheet = sheets[0]
    rows = sheet.get('rows', [])
    headers = [str(h).strip() for h in sheet.get('headers', [])]

    # 自动检测计算类型
    name = card.get('name', '')
    detected = auto_detect_calc_type(name)

    if detected == 'count_rows':
        return len(rows)

    # 年龄相关计算 - 需要找到身份证列或年龄列
    id_col = None
    age_col = None
    for i, h in enumerate(headers):
        h_stripped = h.replace('\n', '').strip()
        if ('身份证' in h_stripped or '身份' in h_stripped or '证件号' in h_stripped) and id_col is None:
            id_col = i
        if ('年龄' in h_stripped or '周岁' in h_stripped) and age_col is None:
            age_col = i

    count = 0
    if detected in ('age_ge_60', 'age_ge_80', 'age_ge_18', 'age_lt_18'):
        for row in rows:
            age = None
            # 优先从身份证号计算年龄
            if id_col is not None and id_col < len(row):
                age = calc_age_from_idcard(row[id_col])
            # 如果没有身份证列，尝试从年龄列读取
            if age is None and age_col is not None and age_col < len(row):
                try:
                    age = int(float(str(row[age_col]).strip()))
                except (ValueError, TypeError):
                    pass
            if age is not None:
                if detected == 'age_ge_60' and age >= 60:
                    count += 1
                elif detected == 'age_ge_80' and age >= 80:
                    count += 1
                elif detected == 'age_ge_18' and age >= 18:
                    count += 1
                elif detected == 'age_lt_18' and age < 18:
                    count += 1
    else:
        count = len(rows)

    return count


# ===== 证明模板定义 =====
CERTIFICATE_TEMPLATES = [
    {
        "id": "cert_residence",
        "icon": "🏠",
        "name": "居住证明",
        "desc": "证明村民在本村常住信息",
        "fields": ["姓名", "身份证号", "性别", "居住地址", "居住时间起", "居住时间止", "用途"]
    },
    {
        "id": "cert_identity",
        "icon": "🪪",
        "name": "身份证明",
        "desc": "证明村民身份信息",
        "fields": ["姓名", "性别", "民族", "出生日期", "身份证号", "住址", "用途"]
    },
    {
        "id": "cert_family_relation",
        "icon": "👨‍👩‍👧‍👦",
        "name": "家庭关系证明",
        "desc": "证明家庭成员关系",
        "fields": ["户主姓名", "户主身份证号", "家庭成员姓名", "与户主关系", "成员身份证号", "用途"]
    },
    {
        "id": "cert_hardship",
        "icon": "🤝",
        "name": "困难证明",
        "desc": "经济困难情况证明",
        "fields": ["姓名", "身份证号", "家庭住址", "家庭人口数", "年收入", "困难原因", "用途"]
    },
    {
        "id": "cert_income",
        "icon": "💰",
        "name": "收入证明",
        "desc": "收入情况证明",
        "fields": ["姓名", "身份证号", "工作单位", "年收入", "月均收入", "用途"]
    },
    {
        "id": "cert_housing",
        "icon": "🏘️",
        "name": "宅基地证明",
        "desc": "宅基地使用情况证明",
        "fields": ["户主姓名", "身份证号", "宅基地位置", "面积(㎡)", "四至界限", "用途"]
    },
    {
        "id": "cert_no_crime",
        "icon": "🛡️",
        "name": "无犯罪记录证明",
        "desc": "无违法犯罪记录证明",
        "fields": ["姓名", "性别", "出生日期", "身份证号", "住址", "用途"]
    },
    {
        "id": "cert_marital",
        "icon": "❤️",
        "name": "婚姻状况证明",
        "desc": "婚姻状态证明",
        "fields": ["姓名", "性别", "身份证号", "婚姻状况", "配偶姓名", "用途"]
    }
]


def get_base_path():
    cfg = load_config()
    return cfg.get('base_path', BASE_DIR)


# ===== 工具函数 =====
def find_item_by_id(cfg, item_id):
    for cat in cfg.get('categories', []):
        for item in cat.get('items', []):
            if item['id'] == item_id:
                return cat, item
            for child in item.get('children', []):
                if child['id'] == item_id:
                    return cat, child
    return None, None


def collect_linked_files(cfg):
    """收集所有已关联到栏目的文件名集合"""
    linked = set()
    for cat in cfg.get('categories', []):
        for item in cat.get('items', []):
            if item.get('file'):
                linked.add(item['file'])
            for child in item.get('children', []):
                if child.get('file'):
                    linked.add(child['file'])
    return linked


def scan_files(base_path):
    """扫描目录中的所有文档文件"""
    if not os.path.isdir(base_path):
        return []
    extensions = [f'*{ext}' for ext in DOCUMENT_EXTENSIONS]
    files = []
    for ext in extensions:
        files.extend(glob.glob(os.path.join(base_path, ext)))
    return sorted({os.path.basename(f) for f in files})


def auto_categorize(base_path, cfg):
    """智能分类：根据文件名关键词自动关联到栏目"""
    all_files = scan_files(base_path)
    # 已关联的文件
    linked_files = collect_linked_files(cfg)

    new_files = [f for f in all_files if f not in linked_files]
    matched = 0
    unmatched = 0

    # 关键词映射：栏目名称 -> 匹配关键词列表（支持模糊匹配）
    keyword_map = {
        '村民户口': ['户口', '户籍', '成员', '名册'],
        '村民花名册': ['花名册', '花名', '手机号', '联系方式'],
        '党员信息': ['党员', '党务', '支部', '党员名单'],
        '残疾人信息': ['残疾', '残联', '残疾人'],
        '低保信息': ['低保', '最低生活保障', '低保户'],
        '集体经济': ['经济', '收入', '资产', '财务', '集体'],
        '农业种植': ['农业', '种植', '早稻', '粮食', '面积', '台账', '耕地'],
        '小麦投保': ['小麦', '投保', '保险', '投保清单'],
        '利剑护蕾': ['护蕾', '关爱', '未成年', '儿童', '重点关爱'],
        '防溺水台账': ['溺水', '水域', '防溺', '巡查'],
        '安置帮教': ['安置', '帮教', '矫正'],
        '矛盾纠纷': ['矛盾', '纠纷', '调解', '排查化解'],
    }

    # 构建栏目名称到栏目的映射，用于后续查找
    cat_item_map = {}  # (cat_index, item_index) -> item
    for ci, cat in enumerate(cfg.get('categories', [])):
        for ii, item in enumerate(cat.get('items', [])):
            cat_item_map[(ci, ii)] = item

    for filename in new_files:
        name_only = os.path.splitext(filename)[0]  # 去掉扩展名
        name_lower = filename.lower()
        name_only_lower = name_only.lower()
        found = False

        # 计算每个空栏目的匹配得分
        best_match = None
        best_score = 0

        for cat in cfg.get('categories', []):
            for item in cat.get('items', []):
                if item.get('file'):
                    continue  # 已关联的跳过
                item_name = item.get('name', '')
                desc = item.get('desc', '')

                # 获取该栏目的关键词列表
                keywords = keyword_map.get(item_name)
                if not keywords:
                    # 回退：用栏目名称和描述作为关键词
                    keywords = list(set([item_name, desc]))

                # 计算匹配分数：关键词在文件名中出现的次数和位置权重
                score = 0
                matched_keywords = []
                for kw in keywords:
                    if not kw:
                        continue
                    kw_lower = kw.lower()
                    # 在完整文件名（含扩展名）中匹配
                    if kw_lower in name_lower:
                        score += len(kw) * 2  # 关键词越长，权重越高
                        matched_keywords.append(kw)
                    # 在纯文件名（无扩展名）中也匹配
                    if kw_lower in name_only_lower:
                        score += len(kw) * 3  # 纯文件名匹配权重更高
                        if kw not in matched_keywords:
                            matched_keywords.append(kw)

                if score > best_score:
                    best_score = score
                    best_match = (item, filename)

        # 如果找到最佳匹配且分数足够高
        if best_match and best_score > 0:
            best_match[0]['file'] = best_match[1]
            matched += 1
            found = True

        if not found:
            unmatched += 1

    # 未归类的文件添加到"未分类资料"
    if unmatched > 0:
        uncategorized = None
        for cat in cfg.get('categories', []):
            if cat.get('id') == 'cat_uncategorized':
                uncategorized = cat
                break
        if not uncategorized:
            uncategorized = {
                "id": "cat_uncategorized",
                "name": "未分类资料",
                "subtitle": "自动扫描 · 智能归类",
                "icon": "📁",
                "items": []
            }
            cfg['categories'].append(uncategorized)

        remaining = [f for f in new_files if not any(
            item.get('file') == f
            for cat in cfg.get('categories', [])
            for item in cat.get('items', [])
        )]
        for f in remaining:
            uncategorized['items'].append({
                "id": "item_auto_" + uuid.uuid4().hex[:8],
                "name": os.path.splitext(f)[0],
                "desc": "自动扫描发现",
                "color": "gray",
                "emoji": "📄",
                "file": f,
                "auto_added": True
            })

    save_config(cfg)
    return {
        "success": True,
        "total_files": len(all_files),
        "matched": matched,
        "unmatched": unmatched,
        "already_linked": len(linked_files),
        "clustered": unmatched > 0
    }


def read_excel_file(filepath):
    """读取Excel文件，返回headers和rows"""
    ext = os.path.splitext(filepath)[1].lower()
    try:
        if ext == '.csv':
            return read_csv_file(filepath)
        elif ext == '.xls':
            return read_xls_file(filepath)
        elif ext == '.xlsx':
            return read_xlsx_file(filepath)
        else:
            return None, None, "不支持的文件格式"
    except Exception as e:
        return None, None, str(e)


def read_xlsx_file(filepath):
    wb = openpyxl.load_workbook(filepath, read_only=True, data_only=True)
    sheets = wb.sheetnames
    result_sheets = []
    for sheet_name in sheets:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            rows.append([str(c) if c is not None else '' for c in row])
        headers = rows[0] if rows else []
        data = rows[1:] if len(rows) > 1 else []
        result_sheets.append({"name": sheet_name, "headers": headers, "rows": data})
    wb.close()
    return result_sheets, None, None


def read_xls_file(filepath):
    wb = xlrd.open_workbook(filepath)
    sheets = wb.sheet_names()
    result_sheets = []
    for sheet_name in sheets:
        ws = wb.sheet_by_name(sheet_name)
        rows = []
        for row_idx in range(ws.nrows):
            row = []
            for col_idx in range(ws.ncols):
                cell = ws.cell(row_idx, col_idx)
                val = cell.value
                if cell.ctype == 3:  # date
                    val = xlrd.xldate_as_datetime(val, wb.datemode).strftime('%Y-%m-%d')
                row.append(str(val) if val != '' else '')
            rows.append(row)
        headers = rows[0] if rows else []
        data = rows[1:] if len(rows) > 1 else []
        result_sheets.append({"name": sheet_name, "headers": headers, "rows": data})
    return result_sheets, None, None


def read_csv_file(filepath):
    encodings = ['utf-8-sig', 'utf-8', 'gbk', 'gb2312', 'gb18030']
    for enc in encodings:
        try:
            with open(filepath, 'r', encoding=enc) as f:
                reader = csv.reader(f)
                rows = [row for row in reader]
            headers = rows[0] if rows else []
            data = rows[1:] if len(rows) > 1 else []
            return [{"name": "Sheet1", "headers": headers, "rows": data}], None, None
        except (UnicodeDecodeError, csv.Error):
            continue
    return None, None, "无法解码CSV文件"


def save_excel_file(filepath, headers, rows, sheet_name='Sheet1'):
    """保存数据到Excel文件"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext == '.csv':
        with open(filepath, 'w', encoding='utf-8-sig', newline='') as f:
            writer = csv.writer(f)
            writer.writerow(headers)
            writer.writerows(rows)
    else:
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = sheet_name
        ws.append(headers)
        for row in rows:
            ws.append(row)
        wb.save(filepath)


def search_in_files(base_path, keyword):
    """在所有文件中搜索关键词"""
    results = []
    cfg = load_config()
    all_files = scan_files(base_path)
    total_matches = 0
    files_with_matches = 0

    for filename in all_files:
        try:
            filepath = resolve_base_file(base_path, filename, must_exist=True)
        except (ValueError, FileNotFoundError):
            continue
        ext = os.path.splitext(filename)[1].lower()
        if ext not in EXCEL_EXTENSIONS:
            continue

        sheets, _, err = read_excel_file(filepath)
        if err or not sheets:
            continue

        file_matches = []
        for sheet in sheets:
            for row_idx, row in enumerate(sheet.get('rows', []), 1):
                row_text = ' '.join(str(c) for c in row)
                if keyword.lower() in row_text.lower():
                    # 找到包含关键词的单元格上下文
                    context_parts = []
                    for cell in row:
                        if keyword.lower() in str(cell).lower():
                            context_parts.append(str(cell))
                    context = ' | '.join(context_parts[:3])
                    file_matches.append({
                        "row": row_idx,
                        "context": context
                    })
                    total_matches += 1

        if file_matches:
            files_with_matches += 1
            # 找到所属分类
            cat_name = "未分类"
            for cat in cfg.get('categories', []):
                for item in cat.get('items', []):
                    if item.get('file') == filename:
                        cat_name = cat.get('name', '未分类')
                        break

            results.append({
                "name": filename,
                "file": filename,
                "cat_name": cat_name,
                "match_count": len(file_matches),
                "matches": file_matches[:10]  # 每个文件最多显示10条
            })

    return {
        "success": True,
        "keyword": keyword,
        "total_files": files_with_matches,
        "total_matches": total_matches,
        "results": results
    }


# ===== 模板定义 =====
TEMPLATES = [
    {
        "id": "tpl_meeting",
        "icon": "📋",
        "name": "村委会议记录",
        "desc": "党支部会议、村委会议等",
        "fields": ["会议名称", "会议日期", "会议地点", "主持人", "记录人", "参会人员", "会议内容", "决议事项"]
    },
    {
        "id": "tpl_visit",
        "icon": "🏠",
        "name": "走访记录",
        "desc": "入户走访、慰问等",
        "fields": ["走访日期", "走访对象", "走访人员", "走访内容", "群众诉求", "处理结果", "备注"]
    },
    {
        "id": "tpl_notice",
        "icon": "📢",
        "name": "通知公告",
        "desc": "村务通知、公示公告等",
        "fields": ["通知标题", "通知日期", "通知内容", "发布人", "备注"]
    },
    {
        "id": "tpl_patrol",
        "icon": "🛡️",
        "name": "巡查记录",
        "desc": "安全巡查、水域巡查等",
        "fields": ["巡查日期", "巡查地点", "巡查人员", "巡查情况", "发现问题", "处理措施", "备注"]
    },
    {
        "id": "tpl_dispute",
        "icon": "⚖️",
        "name": "矛盾纠纷调解",
        "desc": "纠纷排查化解记录",
        "fields": ["纠纷事由", "当事人", "调解日期", "调解人员", "协议内容", "调解结果", "备注"]
    },
    {
        "id": "tpl_summary",
        "icon": "📊",
        "name": "工作总结",
        "desc": "年度/季度工作总结",
        "fields": ["总结标题", "起止日期", "工作概况", "主要成绩", "存在问题", "下一步计划"]
    }
]


def get_reminders_data(base_path, cfg):
    """扫描到期提醒"""
    reminders = []
    danger = warning = info = 0
    today = datetime.now()

    for cat in cfg.get('categories', []):
        for item in cat.get('items', []):
            if not item.get('file'):
                continue
            try:
                filepath = resolve_base_file(base_path, item['file'], must_exist=True)
            except (ValueError, FileNotFoundError):
                continue

            ext = os.path.splitext(item['file'])[1].lower()
            if ext not in EXCEL_EXTENSIONS:
                continue

            sheets, _, _ = read_excel_file(filepath)
            if not sheets:
                continue

            for sheet in sheets:
                headers = [str(h).strip() for h in sheet.get('headers', [])]
                for row in sheet.get('rows', []):
                    for col_idx, cell in enumerate(row):
                        cell_str = str(cell).strip()
                        # 检测日期字段
                        date_val = None
                        for fmt in ['%Y-%m-%d', '%Y/%m/%d', '%Y年%m月%d日']:
                            try:
                                date_val = datetime.strptime(cell_str, fmt)
                                break
                            except ValueError:
                                continue
                        if date_val and date_val > today:
                            days_left = (date_val - today).days
                            header_name = headers[col_idx] if col_idx < len(headers) else '日期'
                            if days_left <= 30:
                                urgency = 'danger'
                                danger += 1
                            elif days_left <= 90:
                                urgency = 'warning'
                                warning += 1
                            elif days_left <= 180:
                                urgency = 'info'
                                info += 1
                            else:
                                continue

                            name_val = ''
                            name_idx = None
                            for i, h in enumerate(headers):
                                if any(kw in h for kw in ['姓名', '名称', '对象', '项目']):
                                    name_idx = i
                                    break
                            if name_idx is not None and name_idx < len(row):
                                name_val = str(row[name_idx])

                            reminders.append({
                                "name": name_val or item.get('name', ''),
                                "source": item.get('name', ''),
                                "date_field": header_name,
                                "date_value": cell_str,
                                "days_left": days_left,
                                "urgency": urgency,
                                "status": f"还剩{days_left}天" if days_left > 0 else "已过期"
                            })

    reminders.sort(key=lambda x: x['days_left'])
    return {
        "success": True,
        "danger": danger,
        "warning": warning,
        "info": info,
        "reminders": reminders[:50]
    }


# ===== 路由 =====

@app.route('/')
def index():
    resp = send_from_directory(BASE_DIR, 'index.html')
    resp.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
    resp.headers['Pragma'] = 'no-cache'
    resp.headers['Expires'] = '0'
    return resp


@app.route('/uploads/<path:filename>')
def serve_upload(filename):
    """提供上传图片的访问"""
    return send_from_directory(UPLOADS_DIR, sanitize_filename(filename))


@app.route('/api/upload_image', methods=['POST'])
def upload_image():
    """接收 base64 图片，保存到 uploads 目录，返回路径"""
    os.makedirs(UPLOADS_DIR, exist_ok=True)
    data = request.get_json(silent=True) or {}
    b64 = data.get('data', '')
    name = re.sub(r'[^0-9A-Za-z_\-\u4e00-\u9fff]+', '_', str(data.get('name', 'img')))[:40] or 'img'
    ext = str(data.get('ext', 'png')).lower().lstrip('.')
    if ext not in ALLOWED_IMAGE_EXTENSIONS:
        return jsonify({"success": False, "message": "不支持的图片格式"}), 400
    if not b64:
        return jsonify({"success": False, "message": "图片数据为空"}), 400
    # 去掉 data URL 前缀
    if ',' in b64:
        b64 = b64.split(',', 1)[1]
    try:
        image_bytes = base64.b64decode(b64, validate=True)
    except Exception:
        return jsonify({"success": False, "message": "图片数据格式无效"}), 400
    if len(image_bytes) > MAX_IMAGE_BYTES:
        return jsonify({"success": False, "message": "图片不能超过5MB"}), 400
    fname = f"{name}_{hashlib.sha256(image_bytes[:1024]).hexdigest()[:12]}.{ext}"
    fpath = os.path.join(UPLOADS_DIR, fname)
    with open(fpath, 'wb') as f:
        f.write(image_bytes)
    return jsonify({"success": True, "url": f"/uploads/{fname}"})


@app.route('/api/call', methods=['POST'])
def api_call():
    data = request.get_json(silent=True) or {}
    method = data.get('method', '')
    args = data.get('args', [])

    handlers = {
        'get_config': handle_get_config,
        'get_login_config': handle_get_login_config,
        'update_login_config': handle_update_login_config,
        "save_remembered_login": handle_save_remembered_login,
        'check_file': handle_check_file,
        'open_file': handle_open_file,
        'read_excel': handle_read_excel,
        'save_excel': handle_save_excel,
        'save_initial_setup': handle_save_initial_setup,
        'update_village_info': handle_update_village_info,
        'add_category': handle_add_category,
        'update_category': handle_update_category,
        'delete_category': handle_delete_category,
        'move_category': handle_move_category,
        'add_item': handle_add_item,
        'update_item': handle_update_item,
        'delete_item': handle_delete_item,
        'move_item': handle_move_item,
        'select_folder': handle_select_folder,
        'set_folder_path': handle_set_folder_path,
        'has_base_path': handle_has_base_path,
        'get_base_path': handle_get_base_path,
        'get_dashboard_stats': handle_get_dashboard_stats,
        'search_files': handle_search_files,
        'get_reminders': handle_get_reminders,
        'get_templates': handle_get_templates,
        'generate_template': handle_generate_template,
        'auto_categorize': handle_auto_categorize,
        'add_sub_item': handle_add_sub_item,
        'delete_sub_item': handle_delete_sub_item,
        'rename_sub_item': handle_rename_sub_item,
        'moveout_sub_item': handle_moveout_sub_item,
        'moveout_main_item': handle_moveout_main_item,
        'export_config': handle_export_config,
        'import_config': handle_import_config,
        'reset_config': handle_reset_config,
        'reset_all_settings': handle_reset_all_settings,
        'reset_basic_settings': handle_reset_basic_settings,
        'has_config': handle_has_config,
        'select_file_for_item': handle_select_file_for_item,
        'scan_new_files': handle_scan_new_files,
        'check_new_files': handle_check_new_files,
        'auto_categorize_files': handle_auto_categorize,
        'reorder_categories': handle_reorder_categories,
        'reorder_items': handle_reorder_items,
        'read_word': handle_read_word,
        'save_word': handle_save_word,
        'remove_sub_from_item': handle_remove_sub_from_item,
        'remove_main_file': handle_remove_main_file,
        'update_sub_item': handle_update_sub_item,
        # 看板配置
        'get_dashboard_config': handle_get_dashboard_config,
        'save_dashboard_config': handle_save_dashboard_config,
        'calc_card_values': handle_calc_card_values,
        'get_excel_files': handle_get_excel_files,
        'calc_chart_data': handle_calc_chart_data,
        'save_tab_order': handle_save_tab_order,
        'get_tab_order': handle_get_tab_order,
        # 60岁到期提醒
        'get_birthday_reminders': handle_get_birthday_reminders,
        'dismiss_reminders': handle_dismiss_reminders,
        'dismiss_all_reminders': handle_dismiss_all_reminders,
        'clear_dismissed_reminders': handle_clear_dismissed_reminders,
        'mark_reminder_read': handle_mark_reminder_read,
        'save_reminder_file': handle_save_reminder_file,
        # 开具证明
        'get_cert_templates': handle_get_cert_templates,
        'get_certificates': handle_get_certificates,
        'add_certificate': handle_add_certificate,
        'update_certificate': handle_update_certificate,
        'delete_certificate': handle_delete_certificate,
        # 账户管理（超级管理员体系）
        'get_accounts': handle_get_accounts,
        'add_account': handle_add_account,
        'update_account': handle_update_account,
        'delete_account': handle_delete_account,
        # 操作日志
        'add_operation_log': handle_add_operation_log,
        'get_operation_logs': handle_get_operation_logs,
        'clear_operation_logs': handle_clear_operation_logs,
        # 登录外观设置
        'update_login_appearance': handle_update_login_appearance,
        # 顶部标题设置
        'update_header_title': handle_update_header_title,
        # 多账户登录验证
        'multi_login': handle_multi_login,
        # 设置窗口标题
        'set_window_title': handle_set_window_title,
    }

    handler = handlers.get(method)
    if handler:
        try:
            result = handler(args)
            return jsonify({"success": True, "result": result})
        except Exception as e:
            return jsonify({"success": False, "message": str(e)})
    else:
        return jsonify({"success": False, "message": f"未知方法: {method}"})


def verify_secret_key(key):
    """验证超级管理员密钥，失败返回 False"""
    login_cfg = load_login_config()
    return key == login_cfg.get('secret_key', 'admin')


# ===== API处理函数 =====

def handle_get_config(args):
    return load_config()


def handle_get_login_config(args):
    return load_login_config()


def handle_update_login_config(args):
    key, user, password, village, new_key = args[0], args[1], args[2], args[3], args[4] if len(args) > 4 else ''
    if not verify_secret_key(key):
        return {"success": False, "message": "密钥验证失败"}
    login_cfg = load_login_config()
    login_cfg['admin_user'] = user
    login_cfg['admin_pass'] = password
    login_cfg['login_village'] = village
    if new_key:
        login_cfg['secret_key'] = new_key
    save_login_config(login_cfg)
    return {"success": True}



def handle_save_remembered_login(args):
    """保存记住的登录信息到 login_config.json"""
    user = args[0] if len(args) > 0 else ""
    remember = args[1] if len(args) > 1 else False
    pwd = args[2] if len(args) > 2 else ""
    login_cfg = load_login_config()
    # 始终保存用户名（方便下次登录）
    login_cfg["remembered_user"] = user
    if remember:
        # 勾选记住密码时才保存密码
        login_cfg["remembered_pass"] = pwd
    else:
        # 不勾选时只清空密码，保留用户名
        login_cfg["remembered_pass"] = ""
    save_login_config(login_cfg)
    return {"success": True}

def handle_check_file(args):
    filename = args[0]
    bp = get_base_path()
    try:
        filepath = resolve_base_file(bp, filename)
        return {"exists": os.path.exists(filepath)}
    except ValueError:
        return {"exists": False}


def handle_open_file(args):
    filename = args[0]
    bp = get_base_path()
    try:
        filepath = resolve_base_file(bp, filename, must_exist=True)
    except (ValueError, FileNotFoundError):
        return {"success": False, "message": "文件不存在或路径无效"}
    try:
        os.startfile(filepath)
    except AttributeError:
        return {"success": False, "message": "当前平台不支持直接打开文件"}
    return {"success": True}


def handle_read_excel(args):
    filename = args[0]
    bp = get_base_path()
    try:
        filepath = resolve_base_file(bp, filename, must_exist=True)
    except (ValueError, FileNotFoundError) as e:
        return {"success": False, "message": str(e)}

    sheets, _, err = read_excel_file(filepath)
    if err:
        return {"success": False, "message": err}

    total_rows = sum(len(s.get('rows', [])) for s in sheets)
    return {
        "success": True,
        "filename": filename,
        "sheets": sheets,
        "total_rows": total_rows
    }


def handle_save_excel(args):
    filename = args[0]
    sheet_data = args[1]  # {"headers": [...], "rows": [...], "sheet_name": "..."}
    bp = get_base_path()
    try:
        filepath = resolve_base_file(bp, filename)
    except ValueError as e:
        return {"success": False, "message": str(e)}

    headers = sheet_data.get('headers', [])
    rows = sheet_data.get('rows', [])
    sheet_name = sheet_data.get('sheet_name', 'Sheet1')

    save_excel_file(filepath, headers, rows, sheet_name)
    return {"success": True}


def handle_save_initial_setup(args):
    village, town, county, preset = args[0], args[1], args[2], args[3]
    cfg = load_config()
    cfg['village_name'] = village
    cfg['town_name'] = town
    cfg['county_name'] = county

    if preset == 'doushan':
        cfg['categories'] = [
            {
                "id": "cat_1",
                "name": "基础民政档案",
                "subtitle": "户籍 · 党员 · 残联 · 值班",
                "icon": "📋",
                "items": [
                    {"id": "item_1", "name": "村民户口", "desc": "户籍档案管理", "color": "red", "emoji": "👨‍👩‍👧‍👦", "file": ""},
                    {"id": "item_2", "name": "村民花名册", "desc": "户籍档案手机号", "color": "red", "emoji": "👥", "file": ""},
                    {"id": "item_3", "name": "党员信息", "desc": "党员花名册", "color": "red", "emoji": "🎖️", "file": ""},
                    {"id": "item_4", "name": "残疾人信息", "desc": "残疾等级评定", "color": "purple", "emoji": "♿", "file": ""},
                    {"id": "item_5", "name": "低保信息", "desc": "低保户信息", "color": "red", "emoji": "📋", "file": ""}
                ]
            },
            {
                "id": "cat_3",
                "name": "其他档案资料",
                "subtitle": "集体经济 · 低保 · 农业 · 水渠",
                "icon": "📂",
                "items": [
                    {"id": "item_10", "name": "集体经济", "desc": "村集体经济", "color": "gold", "emoji": "💰", "file": ""},
                    {"id": "item_12", "name": "农业种植", "desc": "早稻面积台账", "color": "green", "emoji": "🌾", "file": ""},
                    {"id": "item_13", "name": "小麦投保", "desc": "小麦投保清单", "color": "green", "emoji": "🌾", "file": ""}
                ]
            },
            {
                "id": "cat_2",
                "name": "安全专项台账",
                "subtitle": "利剑护蕾 · 防溺水 · 安置帮教 · 矛盾纠纷",
                "icon": "🛡️",
                "items": [
                    {"id": "item_6", "name": "利剑护蕾", "desc": "重点关爱对象台账", "color": "orange", "emoji": "🛡️", "badge": "重点", "file": ""},
                    {"id": "item_7", "name": "防溺水台账", "desc": "水域巡查记录", "color": "orange", "emoji": "🏊", "badge": "重点", "file": ""},
                    {"id": "item_8", "name": "安置帮教", "desc": "安置帮教记录", "color": "teal", "emoji": "🤝", "file": ""},
                    {"id": "item_9", "name": "矛盾纠纷", "desc": "排查化解台账", "color": "teal", "emoji": "⚖️", "file": ""}
                ]
            }
        ]
    else:
        cfg['categories'] = [
            {
                "id": "cat_1",
                "name": "常用资料",
                "subtitle": "基础档案",
                "icon": "📋",
                "items": []
            }
        ]

    # 自动分类
    bp = cfg.get('base_path', BASE_DIR)
    if os.path.exists(bp):
        ac = auto_categorize(bp, cfg)
    else:
        ac = {"success": False}

    save_config(cfg)
    return {"success": True, "auto_categorize": ac}


def handle_update_village_info(args):
    village, town, county = args[0], args[1], args[2]
    cfg = load_config()
    cfg['village_name'] = village
    cfg['town_name'] = town
    cfg['county_name'] = county
    save_config(cfg)
    return {"success": True}


def handle_add_category(args):
    name, subtitle, icon = args[0], args[1], args[2]
    cfg = load_config()
    new_cat = {
        "id": "cat_" + uuid.uuid4().hex[:8],
        "name": name,
        "subtitle": subtitle,
        "icon": icon,
        "items": []
    }
    cfg['categories'].append(new_cat)
    save_config(cfg)
    return {"success": True}


def handle_update_category(args):
    cat_id, name, subtitle, icon = args[0], args[1], args[2], args[3]
    cfg = load_config()
    for cat in cfg.get('categories', []):
        if cat['id'] == cat_id:
            cat['name'] = name
            cat['subtitle'] = subtitle
            cat['icon'] = icon
            break
    save_config(cfg)
    return {"success": True}


def handle_delete_category(args):
    cat_id = args[0]
    cfg = load_config()
    cfg['categories'] = [c for c in cfg.get('categories', []) if c['id'] != cat_id]
    save_config(cfg)
    return {"success": True}


def handle_move_category(args):
    cat_id, direction = args[0], args[1]
    cfg = load_config()
    cats = cfg.get('categories', [])
    idx = next((i for i, c in enumerate(cats) if c['id'] == cat_id), -1)
    if idx == -1:
        return {"success": False}
    if direction == 'up' and idx > 0:
        cats[idx], cats[idx-1] = cats[idx-1], cats[idx]
    elif direction == 'down' and idx < len(cats) - 1:
        cats[idx], cats[idx+1] = cats[idx+1], cats[idx]
    save_config(cfg)
    return {"success": True}


def handle_add_item(args):
    cat_id = args[0]
    # 支持对象格式：add_item(catId, {name, desc, emoji, color, file, badge})
    if len(args) > 1 and isinstance(args[1], dict):
        d = args[1]
        name  = d.get('name', '新栏目')
        desc  = d.get('desc', '')
        emoji = d.get('emoji', '📋')
        color = d.get('color', 'red')
        badge = d.get('badge', '')
        file  = d.get('file', '')
    else:
        name  = args[1] if len(args) > 1 else '新栏目'
        desc  = args[2] if len(args) > 2 else ''
        emoji = args[3] if len(args) > 3 else '📋'
        color = args[4] if len(args) > 4 else 'red'
        badge = args[5] if len(args) > 5 else ''
        file  = args[6] if len(args) > 6 else ''

    cfg = load_config()
    for cat in cfg.get('categories', []):
        if cat['id'] == cat_id:
            new_item = {
                "id": "item_" + uuid.uuid4().hex[:8],
                "name": name,
                "desc": desc,
                "emoji": emoji,
                "color": color,
                "file": file
            }
            if badge:
                new_item['badge'] = badge
            cat['items'].append(new_item)
            break
    save_config(cfg)
    return {"success": True}


def handle_update_item(args):
    item_id = args[0]
    updates = args[1] if len(args) > 1 else {}
    cfg = load_config()
    _, item = find_item_by_id(cfg, item_id)
    if item:
        for k, v in updates.items():
            item[k] = v
        save_config(cfg)
        return {"success": True}
    return {"success": False, "message": "项目不存在"}


def handle_delete_item(args):
    item_id = args[0]
    cfg = load_config()
    for cat in cfg.get('categories', []):
        cat['items'] = [i for i in cat.get('items', []) if i['id'] != item_id]
        # Also remove from children
        for item in cat.get('items', []):
            if 'children' in item:
                item['children'] = [c for c in item['children'] if c['id'] != item_id]
    save_config(cfg)
    return {"success": True}


def handle_move_item(args):
    cat_id, item_id, direction = args[0], args[1], args[2]
    cfg = load_config()
    for cat in cfg.get('categories', []):
        if cat['id'] == cat_id:
            items = cat.get('items', [])
            idx = next((i for i, it in enumerate(items) if it['id'] == item_id), -1)
            if idx == -1:
                break
            if direction == 'up' and idx > 0:
                items[idx], items[idx-1] = items[idx-1], items[idx]
            elif direction == 'down' and idx < len(items) - 1:
                items[idx], items[idx+1] = items[idx+1], items[idx]
            break
    save_config(cfg)
    return {"success": True}


def handle_select_folder(args):
    # 在Web模式下，直接使用当前base_path
    cfg = load_config()
    bp = cfg.get('base_path', BASE_DIR)
    if os.path.exists(bp):
        ac = auto_categorize(bp, cfg)
        save_config(cfg)
        return {"success": True, "path": bp, "auto_categorize": ac}
    return {"success": False, "message": "文件夹不存在"}


def handle_set_folder_path(args):
    """设置新的资料文件夹路径（由前端原生对话框选择后调用）"""
    path = args.get('path', '')
    if not path:
        return {"success": False, "message": "路径为空"}
    if not os.path.exists(path):
        return {"success": False, "message": f"文件夹不存在: {path}"}
    if not os.path.isdir(path):
        return {"success": False, "message": "请选择文件夹，而不是文件"}
    cfg = load_config()
    cfg['base_path'] = path
    save_config(cfg)
    ac = auto_categorize(path, cfg)
    save_config(cfg)
    return {"success": True, "path": path, "auto_categorize": ac}


def handle_has_base_path(args):
    cfg = load_config()
    return bool(cfg.get('base_path'))


def handle_get_base_path(args):
    cfg = load_config()
    return cfg.get('base_path', BASE_DIR)


def handle_get_dashboard_stats(args):
    cfg = load_config()
    bp = cfg.get('base_path', BASE_DIR)
    all_files = scan_files(bp) if os.path.exists(bp) else []

    linked = 0
    cat_stats = []
    for cat in cfg.get('categories', []):
        items_count = 0
        subs_count = 0
        for item in cat.get('items', []):
            items_count += 1
            if item.get('file') and item['file'] in all_files:
                linked += 1
            subs_count += len(item.get('children', []))
        cat_stats.append({
            "name": cat.get('name', ''),
            "icon": cat.get('icon', '📂'),
            "items": items_count,
            "subs": subs_count
        })

    excel_count = sum(1 for f in all_files if f.endswith(('.xls', '.xlsx', '.csv')))
    doc_count = sum(1 for f in all_files if f.endswith(('.doc', '.docx')))

    return {
        "success": True,
        "overview": {
            "total_files": len(all_files),
            "excel_files": excel_count,
            "doc_files": doc_count,
            "linked_files": linked,
            "scanned_files": len(all_files) - linked
        },
        "cat_stats": cat_stats
    }


def handle_search_files(args):
    keyword = args[0]
    bp = get_base_path()
    return search_in_files(bp, keyword)


def handle_get_reminders(args):
    cfg = load_config()
    bp = cfg.get('base_path', BASE_DIR)
    return get_reminders_data(bp, cfg)


def extract_birthday_from_idcard(id_str):
    """从身份证号码提取生日"""
    id_str = str(id_str).strip()
    if len(id_str) == 18:
        try:
            return datetime(int(id_str[6:10]), int(id_str[10:12]), int(id_str[12:14]))
        except (ValueError, IndexError):
            pass
    elif len(id_str) == 15:
        try:
            return datetime(int('19' + id_str[6:8]), int(id_str[8:10]), int(id_str[10:12]))
        except (ValueError, IndexError):
            pass
    return None


def handle_get_birthday_reminders(args):
    """扫描Excel中即将满60岁的人员提醒（提前15天开始提醒，过期30天自动删除）"""
    config = load_dashboard_config()
    bp = get_base_path()
    today = datetime.now()
    # 提前15天开始提醒
    remind_days_ahead = 15
    # 过期超过30天的自动删除不显示
    overdue_expire_days = 30
    reminders_list = []

    filename = args[0] if len(args) > 0 else ''
    if not filename:
        filename = config.get('analysis_file', '') or ''
        # 也尝试从reminder_file获取
        if not filename:
            filename = config.get('reminder_file', '') or ''

    if not filename:
        return {"success": False, "message": "未关联Excel文件，请先在提醒设置中绑定文件"}

    try:
        filepath = resolve_base_file(bp, filename, must_exist=True)
    except (ValueError, FileNotFoundError):
        return {"success": False, "message": "关联文件不存在"}

    sheets, _, err = read_excel_file(filepath)
    if err or not sheets:
        return {"success": False, "message": err or "读取失败"}

    sheet = sheets[0]
    rows = sheet.get('rows', [])
    headers = [str(h).strip() for h in sheet.get('headers', [])]

    # 找列索引 - 更灵活的匹配
    name_col = None
    id_col = None
    phone_col = None
    group_col = None

    for i, h in enumerate(headers):
        h_stripped = h.replace('\n', '').strip()
        # 姓名列
        if name_col is None:
            if '姓名' in h_stripped:
                if '成员' in h_stripped:
                    name_col = i
                elif name_col is None or '户主' not in h_stripped:
                    name_col = i
        # 身份证列
        if id_col is None and ('身份证' in h_stripped or '证件号' in h_stripped or '身份' in h_stripped):
            id_col = i
        # 手机/电话列
        if phone_col is None and ('手机' in h_stripped or '电话' in h_stripped or '联系' in h_stripped):
            phone_col = i
        # 分组列
        if group_col is None and ('组' in h_stripped or '村组' in h_stripped or '组别' in h_stripped):
            group_col = i

    # 如果没找到手机列，尝试"备注"列
    remark_col = None
    if phone_col is None:
        for i, h in enumerate(headers):
            h_stripped = h.replace('\n', '').strip()
            if '备注' in h_stripped or '说明' in h_stripped:
                remark_col = i
                break

    total = 0
    unread = 0
    read_count = 0
    overdue = 0

    phone_pattern = re.compile(r'1[3-9]\d{9}')

    for row_idx, row in enumerate(rows):
        name_val = ''
        phone_val = ''
        birthday = None

        if name_col is not None and name_col < len(row):
            name_val = str(row[name_col]).strip()

        if id_col is not None and id_col < len(row):
            birthday = extract_birthday_from_idcard(row[id_col])

        # 手机号检测
        if phone_col is not None and phone_col < len(row):
            p = str(row[phone_col]).strip()
            if p and p != 'nan' and p != '无' and p != 'None':
                p_clean = p.replace('.0', '').replace('.00', '')
                m = phone_pattern.search(p_clean)
                if m:
                    phone_val = m.group()
                elif len(p_clean) == 11 and p_clean.startswith('1'):
                    phone_val = p_clean

        if not phone_val and remark_col is not None and remark_col < len(row):
            r_val = str(row[remark_col]).strip()
            if r_val and r_val != 'nan' and r_val != 'None':
                r_clean = r_val.replace('.0', '').replace('.00', '')
                m = phone_pattern.search(r_clean)
                if m:
                    phone_val = m.group()
                elif len(r_clean) == 11 and r_clean.startswith('1'):
                    phone_val = r_clean

        if not birthday:
            continue

        total += 1

        # 计算满60岁的日期
        sixty_birthday_year = birthday.year + 60
        try:
            sixty_date = datetime(sixty_birthday_year, birthday.month, birthday.day)
        except ValueError:
            max_day = calendar.monthrange(sixty_birthday_year, birthday.month)[1]
            day = min(birthday.day, max_day)
            sixty_date = datetime(sixty_birthday_year, birthday.month, day)

        days_left = (sixty_date - today).days

        # 只提醒即将到期的（提前15天内），过期超过30天的自动删除
        if days_left < -overdue_expire_days:  # 过期超过30天的不显示
            continue
        if days_left > remind_days_ahead:  # 还很远的也不在未读里
            pass  # 但仍计入total后的read

        status = ''
        urgency_class = ''
        category = 'age60'

        if days_left < 0:
            status = f'已过期{abs(days_left)}天'
            urgency_class = 'overdue'
            overdue += 1
            unread += 1
        elif days_left == 0:
            status = '今天满60岁！'
            urgency_class = 'urgent'
            unread += 1
        elif days_left <= 7:
            status = f'还有{days_left}天满60岁'
            urgency_class = 'urgent'
            unread += 1
        elif days_left <= 30:
            status = f'还有{days_left}天满60岁'
            urgency_class = 'upcoming'
            unread += 1
        else:
            status = f'{sixty_date.strftime("%Y-%m-%d")}满60岁'
            urgency_class = 'normal'
            read_count += 1

        if days_left <= remind_days_ahead or days_left < 0:
            group_val = ''
            if group_col is not None and group_col < len(row):
                group_val = str(row[group_col]).strip()

            title = name_val
            # 提醒内容格式：XX同志于XX年XX月XX日即将年满60岁...
            detail_sixty = sixty_date.strftime('%Y年%m月%d日')
            detail = f'{name_val}同志于{detail_sixty}即将年满60岁，请及时通知该同志办理社保、养老、高龄补贴等相关认证手续。'
            if phone_val:
                detail += f'联系电话：{phone_val}'
            if group_val:
                detail += f'（{group_val}）'

            # 使用身份证号生成稳定的唯一ID，避免Python id()内存地址重启后变化
            id_raw = str(row[id_col]).strip() if id_col is not None and id_col < len(row) else ''
            rid = f"age60_{name_val}_{id_raw}" if id_raw else f"age60_{name_val}_{row_idx}"
            reminders_list.append({
                "id": rid,
                "title": title,
                "detail": detail,
                "phone": phone_val,
                "has_phone": bool(phone_val),
                "birthday_date": sixty_date.strftime('%Y-%m-%d'),
                "days_left": days_left,
                "status": status,
                "urgency": urgency_class,
                "category": category,
                "group": group_val,
                "created_at": today.strftime('%Y-%m-%d %H:%M:%S'),
                "is_read": days_left > remind_days_ahead
            })

    # 按剩余天数排序
    reminders_list.sort(key=lambda x: x['days_left'])

    # 过滤已删除/已忽略的提醒
    dismissed_data = load_dismissed_reminders()
    dismissed_ids = set(dismissed_data.get('dismissed_ids', []))
    if '__all__' in dismissed_ids:
        # 全部清除过
        reminders_list = []
    else:
        reminders_list = [r for r in reminders_list if r.get('id') not in dismissed_ids]

    # 更新已读状态（用户手动标记）
    user_read_ids = set(dismissed_data.get('user_read_ids', []))
    for r in reminders_list:
        if r.get('id') in user_read_ids:
            r['is_read'] = True

    # 重新计算统计（过滤后）
    unread = sum(1 for r in reminders_list if not r.get('is_read'))
    read_count = sum(1 for r in reminders_list if r.get('is_read'))
    overdue = sum(1 for r in reminders_list if r.get('urgency') == 'overdue')

    return {
        "success": True,
        "total": total,
        "unread": unread,
        "read": read_count,
        "overdue": overdue,
        "reminders": reminders_list[:500],
        "filename": filename
    }


def load_dismissed_reminders():
    """加载已忽略/已删除的提醒ID列表"""
    default = {"dismissed_ids": [], "dismissed_at": "", "cleared_all": False, "user_read_ids": []}
    data = load_json_file(DISMISSED_REMINDERS_PATH, default)
    if isinstance(data, list):
        data = {"dismissed_ids": data, "dismissed_at": "", "cleared_all": False, "user_read_ids": []}
    for key in ('dismissed_ids', 'user_read_ids', 'cleared_all'):
        if key not in data:
            data[key] = default[key]
    return data


def save_dismissed_reminders(data):
    save_json_file(DISMISSED_REMINDERS_PATH, data)


def handle_dismiss_reminders(args):
    """删除/忽略指定的提醒（支持单个和批量）"""
    ids_to_dismiss = args[0] if len(args) > 0 else []
    if not ids_to_dismiss:
        return {"success": True, "dismissed_count": 0, "total_dismissed": 0}
    data = load_dismissed_reminders()
    # 如果之前是"全部清除"状态，先重置为普通列表
    if data.get('cleared_all'):
        data['dismissed_ids'] = []
        data['cleared_all'] = False
    existing = set(data.get('dismissed_ids', []))
    for rid in ids_to_dismiss:
        existing.add(str(rid))
    data['dismissed_ids'] = list(existing)
    data['dismissed_at'] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    data['cleared_all'] = False
    save_dismissed_reminders(data)
    return {"success": True, "dismissed_count": len(ids_to_dismiss), "total_dismissed": len(existing)}


def handle_dismiss_all_reminders(args):
    """一键清除所有提醒（全部标记为已读/忽略）"""
    all_ids = args[0] if len(args) > 0 else []
    now = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    if all_ids:
        # 批量删除指定ID，不设置为"全部清除"
        data = load_dismissed_reminders()
        existing = set(data.get('dismissed_ids', []))
        for rid in all_ids:
            existing.add(str(rid))
        data['dismissed_ids'] = list(existing)
        data['dismissed_at'] = now
        data['cleared_all'] = False
    else:
        # 全部清除
        data = {
            "dismissed_ids": ["__all__"],
            "dismissed_at": now,
            "cleared_all": True
        }
    save_dismissed_reminders(data)
    return {"success": True, "cleared_count": len(all_ids) if all_ids else "all", "cleared_at": now}


def handle_clear_dismissed_reminders(args):
    """清空已删除列表（重新显示所有提醒）"""
    data = {"dismissed_ids": [], "dismissed_at": "", "cleared_all": False, "user_read_ids": []}
    save_dismissed_reminders(data)
    return {"success": True}


def handle_mark_reminder_read(args):
    """标记提醒为已读（用户手动标记，不影响屏蔽列表）"""
    ids_to_mark = args[0] if len(args) > 0 else []
    if not ids_to_mark:
        return {"success": True, "marked_count": 0}
    data = load_dismissed_reminders()
    existing = set(data.get('user_read_ids', []))
    for rid in ids_to_mark:
        existing.add(str(rid))
    data['user_read_ids'] = list(existing)
    data['dismissed_at'] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    save_dismissed_reminders(data)
    return {"success": True, "marked_count": len(ids_to_mark)}


def handle_save_reminder_file(args):
    """保存提醒模块关联的Excel文件"""
    config = load_dashboard_config()
    filename = args[0] if len(args) > 0 else ''
    config['reminder_file'] = filename
    # 如果没有analysis_file，也同步设置
    if not config.get('analysis_file') and filename:
        config['analysis_file'] = filename
    save_dashboard_config(config)
    return {"success": True, "filename": filename}


def handle_get_templates(args):
    return {"success": True, "templates": TEMPLATES}


def handle_generate_template(args):
    tpl_id = args[0]
    field_values = args[1]
    tpl = next((t for t in TEMPLATES if t['id'] == tpl_id), None)
    if not tpl:
        return {"success": False, "message": "模板不存在"}

    # 生成简单的文本记录
    bp = get_base_path()
    filename = f"{tpl['name']}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
    filepath = resolve_base_file(bp, filename)

    content = f"{'='*40}\n{tpl['name']}\n{'='*40}\n\n"
    for field in tpl['fields']:
        value = field_values.get(field, '')
        content += f"【{field}】{value}\n\n"
    content += f"\n生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"

    with open(filepath, 'w', encoding='utf-8') as f:
        f.write(content)

    return {"success": True, "message": f"文档已生成：{filename}"}


def handle_auto_categorize(args):
    cfg = load_config()
    bp = cfg.get('base_path', BASE_DIR)
    if not os.path.exists(bp):
        return {"success": False, "message": "文件夹不存在"}
    result = auto_categorize(bp, cfg)
    return result


def handle_add_sub_item(args):
    parent_id, name, file = args[0], args[1], args[2] if len(args) > 2 else ''
    cfg = load_config()
    for cat in cfg.get('categories', []):
        for item in cat.get('items', []):
            if item['id'] == parent_id:
                if 'children' not in item:
                    item['children'] = []
                item['children'].append({
                    "id": "sub_" + uuid.uuid4().hex[:8],
                    "name": name,
                    "file": file,
                    "emoji": "📄"
                })
                save_config(cfg)
                return {"success": True}
    return {"success": False, "message": "父项目不存在"}


def handle_delete_sub_item(args):
    sub_id = args[0]
    cfg = load_config()
    for cat in cfg.get('categories', []):
        for item in cat.get('items', []):
            if 'children' in item:
                item['children'] = [c for c in item['children'] if c['id'] != sub_id]
    save_config(cfg)
    return {"success": True}


def handle_rename_sub_item(args):
    sub_id, new_name = args[0], args[1]
    cfg = load_config()
    for cat in cfg.get('categories', []):
        for item in cat.get('items', []):
            if 'children' in item:
                for child in item['children']:
                    if child['id'] == sub_id:
                        child['name'] = new_name
                        save_config(cfg)
                        return {"success": True}
    return {"success": False}


def handle_moveout_sub_item(args):
    sub_id = args[0]
    cfg = load_config()
    for cat in cfg.get('categories', []):
        for item in cat.get('items', []):
            if 'children' in item:
                for i, child in enumerate(item['children']):
                    if child['id'] == sub_id:
                        # 移到未分类
                        uncategorized = None
                        for c in cfg.get('categories', []):
                            if c.get('id') == 'cat_uncategorized':
                                uncategorized = c
                                break
                        if not uncategorized:
                            uncategorized = {
                                "id": "cat_uncategorized",
                                "name": "未分类资料",
                                "subtitle": "从其他栏目移出",
                                "icon": "📁",
                                "items": []
                            }
                            cfg['categories'].append(uncategorized)

                        uncategorized['items'].append({
                            "id": child['id'],
                            "name": child.get('name', '移出项'),
                            "desc": "从其他栏目移出",
                            "color": "gray",
                            "emoji": child.get('emoji', '📄'),
                            "file": child.get('file', '')
                        })
                        item['children'].pop(i)
                        save_config(cfg)
                        return {"success": True}
    return {"success": False}


def handle_moveout_main_item(args):
    item_id = args[0]
    cfg = load_config()
    for cat in cfg.get('categories', []):
        for i, item in enumerate(cat.get('items', [])):
            if item['id'] == item_id and item.get('file'):
                # 移出主文档到子栏目
                if 'children' not in item:
                    item['children'] = []
                item['children'].append({
                    "id": "sub_" + uuid.uuid4().hex[:8],
                    "name": os.path.splitext(item['file'])[0],
                    "file": item['file'],
                    "emoji": "📄"
                })
                item['file'] = ''
                save_config(cfg)
                return {"success": True}
    return {"success": False}


def handle_export_config(args):
    cfg = load_config()
    return {"success": True, "config": cfg}


def handle_import_config(args):
    new_cfg = args[0]
    save_config(new_cfg)
    return {"success": True}


def handle_reset_config(args):
    save_config(dict(DEFAULT_CONFIG))
    return {"success": True}


def handle_reset_all_settings(args):
    """重置所有设置为初始状态（需验证超级密钥）"""
    key = args[0] if len(args) > 0 else ''
    if not verify_secret_key(key):
        return {"success": False, "message": "超级密钥验证失败"}
    login_cfg = load_login_config()
    # 重置 config.json
    save_config(dict(DEFAULT_CONFIG))
    # 重置 login_config.json：保留 secret_key 和账号密码，清空外观/标题/图标/背景
    reset_login = {
        "admin_user": login_cfg.get("admin_user", "admin"),
        "admin_pass": login_cfg.get("admin_pass", "123456"),
        "secret_key": login_cfg.get("secret_key", "admin"),
        "login_village": "",
        "remembered_user": "",
        "remembered_pass": "",
    }
    reset_login.update(DEFAULT_LOGIN_APPEARANCE)
    save_login_config(reset_login)
    # 重置 dashboard_config.json
    if os.path.exists(DASHBOARD_CONFIG_PATH):
        os.remove(DASHBOARD_CONFIG_PATH)
    # 重置 dismissed_reminders.json
    if os.path.exists(DISMISSED_REMINDERS_PATH):
        os.remove(DISMISSED_REMINDERS_PATH)
    # 重置 accounts.json（保留 super 账户）
    accts = load_accounts()
    if isinstance(accts, dict):
        acct_list = accts.get("accounts", [])
    else:
        acct_list = accts
    super_accts = [a for a in acct_list if a.get('role') == 'super']
    if not super_accts:
        super_accts = [{"id": "super_001", "username": "admin", "password_hash": hash_password("123456"),
                        "role": "super", "display_name": "超级管理员", "enabled": True}]
    save_accounts(super_accts)
    return {"success": True}


def handle_reset_basic_settings(args):
    """系统设置中的重置：只清村名信息、登录页配置、外观设置、顶部标题（不清账户）"""
    key = args[0] if len(args) > 0 else ''
    if not verify_secret_key(key):
        return {"success": False, "message": "超级密钥验证失败"}
    login_cfg = load_login_config()
    # 重置 config.json
    save_config(dict(DEFAULT_CONFIG))
    # 重置 login_config.json：保留 secret_key、账号密码，清空外观/标题/图标/背景
    reset_login = {
        "admin_user": login_cfg.get("admin_user", "admin"),
        "admin_pass": login_cfg.get("admin_pass", "123456"),
        "secret_key": login_cfg.get("secret_key", "admin"),
        "login_village": "",
        "remembered_user": "",
        "remembered_pass": "",
    }
    reset_login.update(DEFAULT_LOGIN_APPEARANCE)
    save_login_config(reset_login)
    return {"success": True}


def handle_select_file_for_item(args):
    """在Web模式下，返回可选择的文件列表"""
    bp = get_base_path()
    files = scan_files(bp) if os.path.exists(bp) else []
    return {"success": True, "files": files}


def handle_scan_new_files(args):
    cfg = load_config()
    bp = cfg.get('base_path', BASE_DIR)
    if not os.path.exists(bp):
        return {"success": False, "message": "文件夹不存在"}

    all_files = scan_files(bp)
    linked = collect_linked_files(cfg)

    new_files = [f for f in all_files if f not in linked]
    return {"success": True, "new_files": new_files, "count": len(new_files)}


def handle_has_config(args):
    """检查是否已有配置文件（有分类数据）"""
    cfg = load_config()
    return bool(cfg.get('village_name') or cfg.get('categories'))


def handle_check_new_files(args):
    """前端check_new_files调用，返回新文件列表"""
    cfg = load_config()
    bp = cfg.get('base_path', BASE_DIR)
    if not os.path.exists(bp):
        return {"success": True, "new_count": 0, "new_files": []}

    all_files = scan_files(bp)
    linked = collect_linked_files(cfg)

    new_files = []
    for f in all_files:
        if f not in linked:
            name_no_ext = os.path.splitext(f)[0]
            new_files.append({"name": f, "name_no_ext": name_no_ext})

    return {"success": True, "new_count": len(new_files), "new_files": new_files}


def handle_reorder_categories(args):
    """拖拽排序分类"""
    new_order = args[0]  # [{id, ...}, ...]
    cfg = load_config()
    old_cats = {c['id']: c for c in cfg.get('categories', [])}
    new_cats = []
    for item in new_order:
        cid = item.get('id', '')
        if cid in old_cats:
            new_cats.append(old_cats[cid])
    # 添加可能遗漏的
    for c in cfg.get('categories', []):
        if c['id'] not in {nc['id'] for nc in new_cats}:
            new_cats.append(c)
    cfg['categories'] = new_cats
    save_config(cfg)
    return {"success": True}


def handle_reorder_items(args):
    """拖拽排序栏目"""
    cat_id = args[0]
    new_order = args[1]  # [{id, ...}, ...]
    cfg = load_config()
    for cat in cfg.get('categories', []):
        if cat['id'] == cat_id:
            old_items = {it['id']: it for it in cat.get('items', [])}
            new_items = []
            for item in new_order:
                iid = item.get('id', '')
                if iid in old_items:
                    new_items.append(old_items[iid])
            for it in cat.get('items', []):
                if it['id'] not in {ni['id'] for ni in new_items}:
                    new_items.append(it)
            cat['items'] = new_items
            break
    save_config(cfg)
    return {"success": True}


def handle_read_word(args):
    """读取Word文档"""
    try:
        from docx import Document
    except ImportError:
        return {"success": False, "message": "python-docx未安装"}

    filename = args[0]
    bp = get_base_path()
    try:
        filepath = resolve_base_file(bp, filename, must_exist=True)
    except (ValueError, FileNotFoundError) as e:
        return {"success": False, "message": str(e)}

    try:
        doc = Document(filepath)
        paragraphs = []
        for p in doc.paragraphs:
            paragraphs.append({
                "text": p.text,
                "style": p.style.name if p.style else "Normal"
            })

        tables = []
        for table in doc.tables:
            rows_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                rows_data.append(row_data)
            tables.append(rows_data)

        return {
            "success": True,
            "filename": filename,
            "paragraphs": paragraphs,
            "tables": tables
        }
    except Exception as e:
        return {"success": False, "message": str(e)}


def handle_save_word(args):
    """保存Word文档"""
    try:
        from docx import Document
    except ImportError:
        return {"success": False, "message": "python-docx未安装"}

    filename = args[0]
    save_data = args[1] if len(args) > 1 else {}
    bp = get_base_path()
    try:
        filepath = resolve_base_file(bp, filename)
    except ValueError as e:
        return {"success": False, "message": str(e)}

    try:
        doc = Document()
        # 添加段落
        for para in save_data.get('paragraphs', []):
            text = para.get('text', '') if isinstance(para, dict) else str(para)
            doc.add_paragraph(text)
        # 添加表格
        for table_data in save_data.get('tables', []):
            if table_data and len(table_data) > 0:
                cols = len(table_data[0]) if table_data[0] else 1
                table = doc.add_table(rows=len(table_data), cols=cols)
                for i, row in enumerate(table_data):
                    for j, cell in enumerate(row):
                        if j < cols:
                            table.rows[i].cells[j].text = str(cell)
        doc.save(filepath)
        return {"success": True}
    except Exception as e:
        return {"success": False, "message": str(e)}


def handle_remove_sub_from_item(args):
    """移除子栏目的文件关联"""
    sub_id = args[0]
    cfg = load_config()
    for cat in cfg.get('categories', []):
        for item in cat.get('items', []):
            if 'children' in item:
                for child in item['children']:
                    if child['id'] == sub_id:
                        child['file'] = ''
                        save_config(cfg)
                        return {"success": True}
    return {"success": False, "message": "子栏目不存在"}


def handle_remove_main_file(args):
    """移除主栏目的文件关联"""
    item_id = args[0]
    cfg = load_config()
    for cat in cfg.get('categories', []):
        for item in cat.get('items', []):
            if item['id'] == item_id:
                item['file'] = ''
                save_config(cfg)
                return {"success": True}
    return {"success": False, "message": "栏目不存在"}


def handle_update_sub_item(args):
    """更新子栏目信息（名称和/或文件）"""
    sub_id = args[0]
    new_name = args[1] if len(args) > 1 else None
    new_file = args[2] if len(args) > 2 else None
    cfg = load_config()
    for cat in cfg.get('categories', []):
        for item in cat.get('items', []):
            if 'children' in item:
                for child in item['children']:
                    if child['id'] == sub_id:
                        if new_name is not None:
                            child['name'] = new_name
                        if new_file is not None:
                            child['file'] = new_file
                        save_config(cfg)
                        return {"success": True}
    return {"success": False, "message": "子栏目不存在"}


# ===== 看板配置 API =====

def handle_get_dashboard_config(args):
    return load_dashboard_config()


def handle_save_dashboard_config(args):
    config = args[0]
    save_dashboard_config(config)
    return {"success": True}


def handle_calc_card_values(args):
    """计算所有看板卡片的数值"""
    config = load_dashboard_config()
    bp = get_base_path()
    results = []
    for card in config.get('cards', []):
        value = calculate_card_value(card, bp)
        detected = auto_detect_calc_type(card.get('name', ''))
        detected_desc = ''
        if detected == 'count_rows':
            detected_desc = '自动统计行数(人数)'
        elif detected == 'age_ge_60':
            detected_desc = '根据身份证号计算≥60岁'
        elif detected == 'age_ge_80':
            detected_desc = '根据身份证号计算≥80岁'
        elif detected == 'age_ge_18':
            detected_desc = '根据身份证号计算≥18岁'
        elif detected == 'age_lt_18':
            detected_desc = '根据身份证号计算<18岁'
        results.append({
            "id": card['id'],
            "name": card.get('name', ''),
            "value": value,
            "icon": card.get('icon', '📊'),
            "color": card.get('color', '#e74c3c'),
            "calc_type": card.get('calc_type', 'auto'),
            "detected": detected,
            "detected_desc": detected_desc
        })
    return {"success": True, "cards": results}


def handle_get_excel_files(args):
    """返回可关联的Excel文件列表"""
    bp = get_base_path()
    files = scan_files(bp) if os.path.exists(bp) else []
    excel_files = [f for f in files if f.endswith(('.xls', '.xlsx', '.csv'))]
    return {"success": True, "files": excel_files}


def handle_calc_chart_data(args):
    """计算数据分析图表数据 - 从关联Excel提取年龄分布和分组数据"""
    config = load_dashboard_config()
    filename = config.get('analysis_file', '')
    if not filename:
        filename = args[0] if len(args) > 0 else ''

    if not filename:
        return {"success": False, "message": "未关联Excel文件"}

    bp = get_base_path()
    try:
        filepath = resolve_base_file(bp, filename, must_exist=True)
    except (ValueError, FileNotFoundError):
        return {"success": False, "message": "文件不存在"}

    sheets, _, err = read_excel_file(filepath)
    if err or not sheets:
        return {"success": False, "message": err or "读取失败"}

    sheet = sheets[0]
    rows = sheet.get('rows', [])
    headers = [str(h).strip() for h in sheet.get('headers', [])]

    # 找身份证列和年龄列
    id_col = None
    age_col = None
    group_col = None
    name_col = None
    for i, h in enumerate(headers):
        h_stripped = h.replace('\n', '').strip()
        if ('身份证' in h_stripped or '身份' in h_stripped or '证件号' in h_stripped) and id_col is None:
            id_col = i
        if ('年龄' in h_stripped or '周岁' in h_stripped) and age_col is None:
            age_col = i
        if ('组' in h_stripped or '组别' in h_stripped) and group_col is None:
            group_col = i
        if '姓名' in h_stripped and name_col is None:
            name_col = i

    # 计算年龄分布
    age_dist = {'0-17岁': 0, '18-35岁': 0, '36-59岁': 0, '60-79岁': 0, '80岁以上': 0}
    for row in rows:
        age = None
        if id_col is not None and id_col < len(row):
            age = calc_age_from_idcard(row[id_col])
        if age is None and age_col is not None and age_col < len(row):
            try:
                age = int(float(str(row[age_col]).strip()))
            except (ValueError, TypeError):
                pass
        if age is not None:
            if age < 18:
                age_dist['0-17岁'] += 1
            elif age <= 35:
                age_dist['18-35岁'] += 1
            elif age <= 59:
                age_dist['36-59岁'] += 1
            elif age <= 79:
                age_dist['60-79岁'] += 1
            else:
                age_dist['80岁以上'] += 1

    # 计算分组人口分布
    group_dist = {}
    if group_col is not None:
        for row in rows:
            if group_col < len(row):
                g = str(row[group_col]).strip()
                if g:
                    group_dist[g] = group_dist.get(g, 0) + 1
    else:
        # 没有分组列时，显示总计
        group_dist = {'总计': len(rows)}

    return {
        "success": True,
        "filename": filename,
        "total_rows": len(rows),
        "age_dist": age_dist,
        "group_dist": group_dist,
        "has_id_col": id_col is not None,
        "has_age_col": age_col is not None,
        "has_group_col": group_col is not None
    }


def handle_save_tab_order(args):
    """保存标签页排序"""
    order = args[0]
    config = load_dashboard_config()
    config['tab_order'] = order
    save_dashboard_config(config)
    return {"success": True}


def handle_get_tab_order(args):
    """获取标签页排序"""
    config = load_dashboard_config()
    return {"success": True, "tab_order": config.get('tab_order', ["files", "dashboard", "search", "reminders", "templates"])}


# ===== 开具证明 API =====

def handle_get_cert_templates(args):
    return {"success": True, "templates": CERTIFICATE_TEMPLATES}


def handle_get_certificates(args):
    data = load_cert_data()
    # 支持按类型筛选
    cert_type = args[0] if len(args) > 0 else None
    certs = data.get('certificates', [])
    if cert_type:
        certs = [c for c in certs if c.get('type_id') == cert_type]
    return {"success": True, "certificates": certs}


def handle_add_certificate(args):
    cert_info = args[0]
    data = load_cert_data()
    cert_id = data.get('next_id', 1)
    now = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    cert = {
        "id": cert_id,
        "code": f"C{cert_id:04d}",
        "type_id": cert_info.get('type_id', ''),
        "type_name": cert_info.get('type_name', ''),
        "applicant": cert_info.get('applicant', ''),
        "fields": cert_info.get('fields', {}),
        "purpose": cert_info.get('purpose', ''),
        "issue_date": now,
        "status": "已开具",
        "issuer": cert_info.get('issuer', ''),
        "created_at": now
    }
    data['certificates'].append(cert)
    data['next_id'] = cert_id + 1
    save_cert_data(data)
    return {"success": True, "certificate": cert}


def handle_update_certificate(args):
    cert_id = int(args[0])
    updates = args[1] if len(args) > 1 else {}
    data = load_cert_data()
    for cert in data.get('certificates', []):
        if cert['id'] == cert_id:
            cert.update(updates)
            save_cert_data(data)
            return {"success": True}
    return {"success": False, "message": "证明记录不存在"}


def handle_delete_certificate(args):
    cert_id = int(args[0])
    data = load_cert_data()
    data['certificates'] = [c for c in data.get('certificates', []) if c['id'] != cert_id]
    save_cert_data(data)
    return {"success": True}


# ===== 账户管理 API =====

def handle_get_accounts(args):
    """获取账户列表（需要超级管理员密钥验证）"""
    key = args[0] if args else ''
    if not verify_secret_key(key):
        return {"success": False, "message": "密钥验证失败"}
    accts = load_accounts()
    # 不返回密码
    safe = []
    acct_list = accts.get('accounts', accts) if isinstance(accts, dict) else (accts if isinstance(accts, list) else [])
    for a in acct_list:
        safe.append({k: v for k, v in a.items() if k not in ('password', 'password_hash')})
    return {"success": True, "accounts": safe}


def handle_add_account(args):
    """新增账户（超级管理员操作）"""
    key, username, password, role, display_name = (
        args[0] if len(args) > 0 else '',
        args[1] if len(args) > 1 else '',
        args[2] if len(args) > 2 else '123456',
        args[3] if len(args) > 3 else 'staff',
        args[4] if len(args) > 4 else '',
    )
    if not verify_secret_key(key):
        return {"success": False, "message": "密钥验证失败"}
    if not username:
        return {"success": False, "message": "账号不能为空"}
    accts = load_accounts()
    # 检查重复
    for a in accts.get('accounts', []):
        if a['username'] == username:
            return {"success": False, "message": "账号已存在"}
    new_id = f"u_{uuid.uuid4().hex[:8]}"
    accts['accounts'].append({
        "id": new_id,
        "username": username,
        "password_hash": hash_password(password),
        "role": role,
        "display_name": display_name or username,
        "created_at": datetime.now().strftime('%Y-%m-%d %H:%M'),
        "enabled": True
    })
    save_accounts(accts)
    return {"success": True, "id": new_id}


def handle_update_account(args):
    """修改账户（超级管理员操作）"""
    key = args[0] if len(args) > 0 else ''
    account_id = args[1] if len(args) > 1 else ''
    updates = args[2] if len(args) > 2 else {}
    if not verify_secret_key(key):
        return {"success": False, "message": "密钥验证失败"}
    accts = load_accounts()
    for a in accts.get('accounts', []):
        if a['id'] == account_id:
            # 不允许修改超级管理员的角色
            if a['role'] == 'super' and updates.get('role') not in (None, 'super'):
                return {"success": False, "message": "不能降级超级管理员"}
            for k, v in updates.items():
                if k == 'password':
                    a['password_hash'] = hash_password(v)
                    a.pop('password', None)
                elif k not in ('id', 'password_hash'):
                    a[k] = v
            save_accounts(accts)
            return {"success": True}
    return {"success": False, "message": "账户不存在"}


def handle_delete_account(args):
    """删除账户（超级管理员操作）"""
    key = args[0] if len(args) > 0 else ''
    account_id = args[1] if len(args) > 1 else ''
    if not verify_secret_key(key):
        return {"success": False, "message": "密钥验证失败"}
    accts = load_accounts()
    before = accts.get('accounts', [])
    # 不允许删除超级管理员
    target = next((a for a in before if a['id'] == account_id), None)
    if target and target['role'] == 'super':
        return {"success": False, "message": "不能删除超级管理员账户"}
    accts['accounts'] = [a for a in before if a['id'] != account_id]
    save_accounts(accts)
    return {"success": True}


# ===== 操作日志 =====

LOGS_PATH = os.path.join(UPLOADS_DIR, 'operation_logs.json')

def _ensure_logs_dir():
    """确保操作日志所在目录存在"""
    os.makedirs(UPLOADS_DIR, exist_ok=True)

def _load_logs():
    data = load_json_file(LOGS_PATH, [])
    return data if isinstance(data, list) else []

def _save_logs(logs):
    # 保留最近500条
    logs = logs[-500:]
    _ensure_logs_dir()
    save_json_file(LOGS_PATH, logs)

def handle_add_operation_log(args):
    """新增一条操作日志"""
    operator = args[0] if len(args) > 0 else '未知'
    action = args[1] if len(args) > 1 else '操作'
    detail = args[2] if len(args) > 2 else ''
    log_entry = {
        "time": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        "operator": operator,
        "action": action,
        "detail": detail
    }
    logs = _load_logs()
    logs.append(log_entry)
    _save_logs(logs)
    return {"success": True}

def handle_get_operation_logs(args):
    """获取操作日志（按时间倒序）"""
    logs = _load_logs()
    # 按时间倒序返回
    logs.reverse()
    return {"logs": logs}

def handle_clear_operation_logs(args):
    """清空所有操作日志"""
    if os.path.exists(LOGS_PATH):
        os.remove(LOGS_PATH)
    return {"success": True}


def handle_update_login_appearance(args):
    """更新登录页外观（图标、背景、标题）"""
    key = args[0] if len(args) > 0 else ''
    updates = args[1] if len(args) > 1 else {}
    if not verify_secret_key(key):
        return {"success": False, "message": "密钥验证失败"}
    login_cfg = load_login_config()
    allowed_keys = {
        'login_title', 'login_subtitle', 'login_footer_text',
        'login_icon_type', 'login_icon_emoji', 'login_icon_image', 'login_icon_crop',
        'login_bg_type', 'login_bg_color', 'login_bg_image', 'login_bg_fit',
        'window_title', 'login_village',
    }
    for k, v in updates.items():
        if k in allowed_keys:
            login_cfg[k] = v
    save_login_config(login_cfg)
    return {"success": True}


def handle_update_header_title(args):
    """更新顶部标题文字"""
    header_title = args[0] if len(args) > 0 else ''
    header_subtitle = args[1] if len(args) > 1 else ''
    key = args[2] if len(args) > 2 else ''
    if not verify_secret_key(key):
        return {"success": False, "message": "密钥验证失败"}
    login_cfg = load_login_config()
    login_cfg['header_title'] = header_title
    login_cfg['header_subtitle'] = header_subtitle
    save_login_config(login_cfg)
    return {"success": True}


def handle_multi_login(args):
    """多账户登录验证（新的登录方式）"""
    username = args[0] if len(args) > 0 else ''
    password = args[1] if len(args) > 1 else ''
    accts = load_accounts()
    for a in accts.get('accounts', []):
        if a.get('username') == username and verify_password(a, password) and a.get('enabled', True):
            if migrate_account_password(a):
                save_accounts(accts)
            return {
                "success": True,
                "id": a.get('id', ''),
                "role": a.get('role', 'staff'),
                "display_name": a.get('display_name', username)
            }
    return {"success": False, "message": "账号或密码错误"}


def handle_set_window_title(args):
    """设置窗口标题（通过pywebview API） - 仅在原生环境中有效"""
    title = args[0] if len(args) > 0 else '智慧文档管理系统'
    try:
        import pywebview
        # pywebview的window对象是全局单例
        for w in pywebview.windows:
            w.set_title(title)
            break
        return {"success": True, "title": title}
    except Exception:
        return {"success": True, "title": title, "note": "非原生环境"}


# ===== 启动 =====
if __name__ == '__main__':
    # 直接运行时（开发模式），由 launcher_webview.py 统一启动
    app.run(host='127.0.0.1', port=0, debug=False)
